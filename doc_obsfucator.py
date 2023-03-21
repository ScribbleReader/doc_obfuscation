from docx import Document
import docx
import random
from string import ascii_uppercase
from string import ascii_lowercase
import os
import re
from PIL import Image
import io
import pickle

def get_pictures(doc, result_path):
    try:
        dict_rel = doc.part._rels
        for rel in dict_rel:
            rel = dict_rel[rel]
            if "image" in rel.target_ref:
                if not os.path.exists(result_path):
                    os.makedirs(result_path)
                img_name = re.findall("/(.*)", rel.target_ref)[0]
                with open(f'{result_path}/{img_name}', "wb") as f:
                    f.write(rel.target_part.blob)
    except:
        pass

def change_doc(doc, mapper):
    for para in doc.paragraphs:
            for run in para.runs:
                for i in range(len(run.text)):
                    if '\u4e00' <= run.text[i] <= '\u9fa5' or ((run.text[i].isalpha() or run.text[i].isdigit())\
                                                                and run.text[i].isascii()):
                        mapped = mapper[run.text[i]]
                        run.text = run.text[:i] + mapped + run.text[i+1:]


    for table in doc.tables:
        for cell in table._cells:
            for i in range(len(cell.text)):
                if '\u4e00' <= cell.text[i] <= '\u9fa5' or ((cell.text[i].isalpha() or cell.text[i].isdigit())\
                                                             and cell.text[i].isascii()):
                    mapped = mapper[cell.text[i]]
                    cell.text = cell.text[:i] + mapped + cell.text[i+1:]
    return doc


def change_img(doc):
    for rel in doc.part._rels:
        val = doc.part._rels[rel]
        if "image" in val.target_ref:
            # print(val.target_part.blob)

            png_image = Image.open(io.BytesIO(val.target_part.blob))
            yellow_image = Image.new('RGB', png_image.size, color='black')
            t = doc.part._rels[rel].target_part

            doc.part._rels[rel].target_part.set_blob(yellow_image.tobytes())
    return doc

def recover_img(doc, img_path):
    for rel in doc.part._rels:
        val = doc.part._rels[rel]
        if "image" in val.target_ref:
            # print(val.target_part.blob)
            img_name = re.findall("/(.*)", val.target_ref)[0]
            img = open(os.path.join(img_path, img_name), 'rb').read()

            doc.part._rels[rel].target_part.set_blob(img)
    return doc

def generate_mapper(doc):
    collecter = set()
    for para in doc.paragraphs:
        for run in para.runs:
            for i in range(len(run.text)):
                if '\u4e00' <= run.text[i] <= '\u9fa5':
                    collecter.add(run.text[i])

    for table in doc.tables:
        for cell in table._cells:
            for i in range(len(cell.text)):
                if '\u4e00' <= cell.text[i] <= '\u9fa5':
                    collecter.add(cell.text[i])

    left = list(collecter)
    right = list(collecter)

    eng_1 = list(ascii_uppercase)
    eng_11 = list(ascii_uppercase)

    eng_2 = list(ascii_lowercase) + [str(i) for i in range(10)]
    eng_22 = list(ascii_lowercase) + [str(i) for i in range(10)]

    random.shuffle(right)
    random.shuffle(eng_11)
    random.shuffle(eng_22)

    mapper = {}
    for l, r in zip(left, right):
        mapper[l] = r

    for l, r in zip(eng_1, eng_11):
        mapper[l] = r

    for l, r in zip(eng_2, eng_22):
        mapper[l] = r

    return mapper


if __name__ == '__main__':
    # 混淆word
    doc = Document('your_doc.docx')

    # 储存图片到指定路径
    get_pictures(doc, 'img_path')

    mapper = generate_mapper(doc)

    doc = change_doc(doc, mapper=mapper)
    doc = change_img(doc)

    # 保存映射关系
    with open('mapper.data', 'wb') as f:
        pickle.dump(mapper, f)

    doc.save('document1.docx')

    # 恢复
    doc2 = Document('document2.docx')

    # 加载映射关系
    mapper = pickle.load(open('mapper.data', 'rb'))
    reverse_mapper = {v: k for k, v in mapper.items()}

    # 恢复文字
    doc2 = change_doc(doc2, reverse_mapper)
    # 恢复图片
    doc2 = recover_img(doc2, img_path='img_path')

    doc2.save('document2.docx')
